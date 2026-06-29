# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# A4 가로 (Landscape)
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# 기본 폰트 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def write_cell(cell, text, bold=False, size=9, fg='000000', center=True, bg=None):
    if bg:
        set_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))

# ── 제목 ──────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BeautyScope  기능 명세서')
run.bold = True
run.font.size = Pt(18)
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x6D)

doc.add_paragraph()

# ── 프로젝트 개요 ───────────────────────────────
for label, val in [
    ('프로젝트명', 'BeautyScope (화장품 리뷰 분석 사이트)'),
    ('개  발  환  경', 'Spring MVC, MyBatis, Oracle, JSP'),
    ('개  발  인  원', '1명'),
    ('특  이  사  항', '쿠팡 웹 크롤링 수집 데이터 활용'),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label} : ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = '맑은 고딕'
    r2 = p.add_run(val)
    r2.font.size = Pt(10)
    r2.font.name = '맑은 고딕'

doc.add_paragraph()

# ── 테이블 ────────────────────────────────────────
HEADERS = ['기능-ID','화면-ID','Category','Depth1','Depth2','Depth3','중요도','지원디바이스','기능 설명']
WIDTHS  = [1.8, 2.5, 2.2, 2.8, 2.8, 2.0, 1.5, 2.2, 8.9]  # 합계 26.7cm

DATA = [
    ('V1-101','Main',          '메인',   'TOP5 랭킹',        '',             '',       '상','웹','평점 기준 상위 5개 제품을 카드 형태로 메인 화면에 노출'),
    ('V1-102','Main',          '메인',   '카테고리 바로가기','',             '',       '중','웹','카테고리별 빠른 접근 링크 제공'),
    ('V1-201','Product List',  '제품관리','제품 목록',        '카테고리 필터','',       '상','웹','스킨케어 / 메이크업 등 카테고리별 필터링'),
    ('V1-202','Product List',  '제품관리','제품 목록',        '정렬',         '평점순', '상','웹','평점 높은순 / 낮은순 정렬'),
    ('V1-203','Product List',  '제품관리','제품 목록',        '정렬',         '최신순', '중','웹','최근 등록순 정렬'),
    ('V1-204','Product List',  '제품관리','제품 목록',        '페이징',       '',       '상','웹','페이지 단위로 목록 표시'),
    ('V1-301','Product Detail','제품관리','제품 상세',        '제품 정보',    '',       '상','웹','제품명, 브랜드, 카테고리, 평균 별점 표시'),
    ('V1-302','Product Detail','제품관리','제품 상세',        '별점 분포',    '차트',   '상','웹','1점~5점 별점 분포 막대 차트 (Chart.js)'),
    ('V1-303','Product Detail','제품관리','제품 상세',        '리뷰 목록',    '비동기', '상','웹','해당 제품의 크롤링 리뷰 목록 비동기(fetch) 로딩'),
    ('V1-401','Search',        '검색',   '키워드 검색',      '제품명',       '',       '상','웹','제품명으로 검색 (LIKE 쿼리)'),
    ('V1-402','Search',        '검색',   '키워드 검색',      '리뷰 내용',    '',       '중','웹','리뷰 내용 키워드로 검색'),
    ('V1-501','Ranking',       '분석',   '카테고리 랭킹',    '',             '',       '상','웹','카테고리별 평점 TOP 제품 랭킹 목록'),
    ('V1-502','Ranking',       '분석',   '브랜드 비교',      '평균 평점',    '차트',   '중','웹','브랜드별 평균 평점 비교 차트 (Chart.js)'),
]

table = doc.add_table(rows=1, cols=len(HEADERS))
table.style = 'Table Grid'

# 헤더
hdr = table.rows[0]
for i, cell in enumerate(hdr.cells):
    cell.width = Cm(WIDTHS[i])
    write_cell(cell, HEADERS[i], bold=True, size=9, fg='FFFFFF', center=True, bg='1F496D')

# 데이터
for ri, row_data in enumerate(DATA):
    bg = 'D6EAF8' if ri % 2 == 0 else 'FFFFFF'
    row = table.add_row()
    for i, cell in enumerate(row.cells):
        cell.width = Cm(WIDTHS[i])
        is_desc = (i == 8)
        write_cell(cell, row_data[i], size=9, fg='000000', center=not is_desc, bg=bg)

doc.save(r'D:\spring-workspace\Mission-Spring\기능명세서.docx')
print('완료')
